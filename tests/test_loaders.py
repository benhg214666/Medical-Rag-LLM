"""文件載入器測試。

全部離線執行：DOCX 用 python-docx 動態建立，PDF 用 pypdf 建立或以 monkeypatch 模擬，
不依賴任何預先準備的測試檔案或網路資源。
"""

from pathlib import Path

import pytest

from app.ingestion.exceptions import (
    DocumentLoadError,
    EmptyDocumentError,
    UnsupportedFileTypeError,
)
from app.ingestion.loaders import load_docx, load_document, load_pdf, load_txt


class TestLoadTxt:
    """純文字檔載入。"""

    def test_chinese_and_english_content(self, tmp_path: Path) -> None:
        file_path = tmp_path / "note.txt"
        content = "病患主訴：胸痛三天。\nPatient reports chest pain for 3 days."
        file_path.write_text(content, encoding="utf-8")

        documents = load_txt(file_path)

        assert len(documents) == 1
        assert documents[0].text == content
        assert "病患主訴" in documents[0].text
        assert "chest pain" in documents[0].text

    def test_bom_is_stripped(self, tmp_path: Path) -> None:
        file_path = tmp_path / "bom.txt"
        # UTF-8 BOM 若未處理，第一個字元會變成不可見的 ﻿，污染後續比對
        file_path.write_bytes("﻿病患資料".encode("utf-8"))

        documents = load_txt(file_path)

        assert not documents[0].text.startswith("﻿")
        assert documents[0].text == "病患資料"

    def test_chinese_not_corrupted(self, tmp_path: Path) -> None:
        file_path = tmp_path / "zh.txt"
        content = "血壓 120/80 mmHg，體溫 36.8°C，給予 Aspirin 100 mg。"
        file_path.write_text(content, encoding="utf-8")

        assert load_txt(file_path)[0].text == content

    def test_metadata_fields(self, tmp_path: Path) -> None:
        file_path = tmp_path / "meta.txt"
        file_path.write_text("內容", encoding="utf-8")

        document = load_txt(file_path)[0]

        assert document.file_name == "meta.txt"
        assert document.file_type == "txt"
        assert document.source == "meta.txt"
        assert document.page_number is None
        assert document.paragraph_number is None

    def test_empty_file_raises(self, tmp_path: Path) -> None:
        file_path = tmp_path / "empty.txt"
        file_path.write_text("", encoding="utf-8")

        with pytest.raises(EmptyDocumentError):
            load_txt(file_path)

    def test_whitespace_only_file_raises(self, tmp_path: Path) -> None:
        file_path = tmp_path / "blank.txt"
        file_path.write_text("   \n\n\t  ", encoding="utf-8")

        with pytest.raises(EmptyDocumentError):
            load_txt(file_path)

    def test_invalid_encoding_raises_clearly(self, tmp_path: Path) -> None:
        file_path = tmp_path / "bad.txt"
        # 不合法的 UTF-8 位元組序列
        file_path.write_bytes(b"\xff\xfe\x00\x81\x82\x83")

        with pytest.raises(DocumentLoadError) as exc_info:
            load_txt(file_path)
        assert "UTF-8" in str(exc_info.value)


class TestLoadDocx:
    """Word 文件載入（動態建立測試檔）。"""

    @staticmethod
    def _create_docx(path: Path, paragraphs: list[str]) -> None:
        import docx

        document = docx.Document()
        for text in paragraphs:
            document.add_paragraph(text)
        document.save(str(path))

    def test_reads_non_empty_paragraphs(self, tmp_path: Path) -> None:
        file_path = tmp_path / "note.docx"
        self._create_docx(file_path, ["第一段：主訴。", "第二段：檢查結果。"])

        documents = load_docx(file_path)

        assert len(documents) == 2
        assert documents[0].text == "第一段：主訴。"
        assert documents[1].text == "第二段：檢查結果。"

    def test_blank_paragraphs_are_skipped(self, tmp_path: Path) -> None:
        file_path = tmp_path / "gaps.docx"
        self._create_docx(file_path, ["有內容", "", "   ", "也有內容"])

        documents = load_docx(file_path)

        assert len(documents) == 2

    def test_paragraph_numbers_reflect_original_position(
        self, tmp_path: Path
    ) -> None:
        file_path = tmp_path / "numbered.docx"
        self._create_docx(file_path, ["第一段", "", "第三段"])

        documents = load_docx(file_path)

        # 空段落被跳過，但編號不前移，才能對應回 Word 中看到的位置
        assert [d.paragraph_number for d in documents] == [1, 3]

    def test_metadata_fields(self, tmp_path: Path) -> None:
        file_path = tmp_path / "meta.docx"
        self._create_docx(file_path, ["內容"])

        document = load_docx(file_path)[0]

        assert document.file_type == "docx"
        assert document.file_name == "meta.docx"
        assert document.page_number is None
        assert document.paragraph_number == 1

    def test_document_with_no_paragraphs_raises(self, tmp_path: Path) -> None:
        file_path = tmp_path / "blank.docx"
        self._create_docx(file_path, ["", "  ", ""])

        with pytest.raises(EmptyDocumentError):
            load_docx(file_path)

    def test_corrupted_file_raises(self, tmp_path: Path) -> None:
        file_path = tmp_path / "broken.docx"
        file_path.write_bytes(b"this is not a valid docx package")

        with pytest.raises(DocumentLoadError):
            load_docx(file_path)


class TestLoadPdf:
    """PDF 載入。

    主要以 monkeypatch 模擬 pypdf，避免為了測試而引入大型 PDF 樣本檔；
    另外用 pypdf 建立真實的空白 PDF，驗證掃描式 PDF 的錯誤訊息。
    """

    @staticmethod
    def _patch_reader(monkeypatch: pytest.MonkeyPatch, page_texts: list[str]) -> None:
        """把 pypdf.PdfReader 換成回傳指定頁面文字的假物件。"""

        class FakePage:
            def __init__(self, text: str) -> None:
                self._text = text

            def extract_text(self) -> str:
                return self._text

        class FakeReader:
            def __init__(self, _path: str) -> None:
                self.pages = [FakePage(text) for text in page_texts]

        import pypdf

        monkeypatch.setattr(pypdf, "PdfReader", FakeReader)

    def test_extracts_text_per_page(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        self._patch_reader(monkeypatch, ["第一頁：主訴。", "第二頁：診斷。"])
        file_path = tmp_path / "report.pdf"
        file_path.write_bytes(b"%PDF-1.4 placeholder")

        documents = load_pdf(file_path)

        assert len(documents) == 2
        assert documents[0].text == "第一頁：主訴。"
        assert documents[1].text == "第二頁：診斷。"

    def test_page_numbers_start_at_one(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        self._patch_reader(monkeypatch, ["A", "B", "C"])
        file_path = tmp_path / "pages.pdf"
        file_path.write_bytes(b"%PDF-1.4 placeholder")

        documents = load_pdf(file_path)

        # 頁碼要與 PDF 閱讀器顯示一致，從 1 起算而非 0
        assert [d.page_number for d in documents] == [1, 2, 3]

    def test_blank_pages_are_skipped_but_numbering_holds(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        self._patch_reader(monkeypatch, ["第一頁", "   ", "第三頁"])
        file_path = tmp_path / "gaps.pdf"
        file_path.write_bytes(b"%PDF-1.4 placeholder")

        documents = load_pdf(file_path)

        assert len(documents) == 2
        assert [d.page_number for d in documents] == [1, 3]

    def test_metadata_includes_total_pages(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        self._patch_reader(monkeypatch, ["A", "B"])
        file_path = tmp_path / "meta.pdf"
        file_path.write_bytes(b"%PDF-1.4 placeholder")

        document = load_pdf(file_path)[0]

        assert document.file_type == "pdf"
        assert document.metadata["total_pages"] == 2
        assert document.paragraph_number is None

    def test_scanned_pdf_reports_ocr_limitation(self, tmp_path: Path) -> None:
        # 用 pypdf 建立真實但無文字圖層的 PDF，模擬掃描件
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        writer.add_blank_page(width=200, height=200)
        file_path = tmp_path / "scanned.pdf"
        with file_path.open("wb") as handle:
            writer.write(handle)

        with pytest.raises(EmptyDocumentError) as exc_info:
            load_pdf(file_path)

        message = str(exc_info.value)
        assert "掃描式" in message
        assert "OCR" in message

    def test_corrupted_pdf_raises(self, tmp_path: Path) -> None:
        file_path = tmp_path / "broken.pdf"
        file_path.write_bytes(b"definitely not a pdf")

        with pytest.raises(DocumentLoadError):
            load_pdf(file_path)


class TestLoadDocumentDispatch:
    """統一入口的分派與驗證。"""

    def test_dispatches_txt(self, tmp_path: Path) -> None:
        file_path = tmp_path / "a.txt"
        file_path.write_text("內容", encoding="utf-8")

        assert load_document(file_path)[0].file_type == "txt"

    def test_dispatches_docx(self, tmp_path: Path) -> None:
        import docx

        document = docx.Document()
        document.add_paragraph("段落內容")
        file_path = tmp_path / "a.docx"
        document.save(str(file_path))

        assert load_document(file_path)[0].file_type == "docx"

    def test_extension_matching_is_case_insensitive(self, tmp_path: Path) -> None:
        # Linux 檔名大小寫敏感，.TXT 也必須能被辨識
        file_path = tmp_path / "upper.TXT"
        file_path.write_text("內容", encoding="utf-8")

        assert load_document(file_path)[0].file_type == "txt"

    @pytest.mark.parametrize("name", ["a.csv", "a.xlsx", "a.json", "a.exe", "noext"])
    def test_unsupported_extension_raises(self, tmp_path: Path, name: str) -> None:
        file_path = tmp_path / name
        file_path.write_text("內容", encoding="utf-8")

        with pytest.raises(UnsupportedFileTypeError):
            load_document(file_path)

    def test_missing_file_raises(self, tmp_path: Path) -> None:
        with pytest.raises(DocumentLoadError):
            load_document(tmp_path / "does_not_exist.txt")

    def test_directory_path_raises(self, tmp_path: Path) -> None:
        directory = tmp_path / "folder.txt"
        directory.mkdir()

        with pytest.raises(DocumentLoadError):
            load_document(directory)
