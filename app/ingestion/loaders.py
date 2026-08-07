"""文件載入器（loaders）。

責任：把不同格式的檔案（TXT / PDF / DOCX）統一轉換成 list[LoadedDocument]，
讓管線後續階段不需要再關心原始格式。這是典型的 adapter 模式：
用一個共同介面隔離掉格式差異。

「文件單位」的切法依格式而異，目的是保留最有用的來源定位資訊：
  - PDF  每頁一個單位，記錄 page_number（頁碼是醫師查證原文最常用的線索）
  - DOCX 每個非空段落一個單位，記錄 paragraph_number
  - TXT  整份檔案一個單位（純文字無內建結構可依循）

安全性：本模組只「讀取」檔案內容，絕不執行上傳內容，也不依 MIME type 決定行為。
"""

import logging
from pathlib import Path

from app.ingestion.exceptions import (
    DocumentLoadError,
    EmptyDocumentError,
    UnsupportedFileTypeError,
)
from app.ingestion.models import LoadedDocument

logger = logging.getLogger(__name__)

# 支援的副檔名（一律小寫比對；Linux 檔名大小寫敏感，故先 lower() 再判斷）
SUPPORTED_EXTENSIONS: frozenset[str] = frozenset({".txt", ".pdf", ".docx"})


def _validate_file(file_path: Path) -> None:
    """檢查檔案可被載入。

    Raises:
        DocumentLoadError: 檔案不存在，或不是一般檔案（例如目錄、裝置檔）。
        UnsupportedFileTypeError: 副檔名不在支援清單內。
    """
    if not file_path.exists():
        raise DocumentLoadError(f"檔案不存在：{file_path.name}")

    if not file_path.is_file():
        raise DocumentLoadError(f"路徑不是一般檔案：{file_path.name}")

    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise UnsupportedFileTypeError(
            f"不支援的檔案格式 '{suffix or '(無副檔名)'}'，目前僅支援：{supported}"
        )


def load_txt(file_path: Path) -> list[LoadedDocument]:
    """載入純文字檔。

    編碼策略：優先 utf-8-sig（能自動去除 BOM），失敗後退回 utf-8。
    兩者皆失敗時視為編碼錯誤，明確拋出例外而不靜默替換成亂碼 —— 在醫療情境下，
    一個被錯誤解碼的字元可能改變藥名或劑量。

    Raises:
        DocumentLoadError: 檔案無法以 UTF-8 系列編碼讀取。
        EmptyDocumentError: 檔案為空或只有空白。
    """
    raw_bytes = file_path.read_bytes()

    text: str | None = None
    for encoding in ("utf-8-sig", "utf-8"):
        try:
            text = raw_bytes.decode(encoding)
            break
        except UnicodeDecodeError:
            continue

    if text is None:
        raise DocumentLoadError(
            f"無法以 UTF-8 解碼檔案 '{file_path.name}'，"
            "請確認檔案編碼（本階段僅支援 UTF-8 / UTF-8-SIG）"
        )

    if not text.strip():
        raise EmptyDocumentError(f"檔案 '{file_path.name}' 沒有任何文字內容")

    return [
        LoadedDocument(
            text=text,
            source=file_path.name,
            file_name=file_path.name,
            file_type="txt",
            page_number=None,
            paragraph_number=None,
            metadata={},
        )
    ]


def load_pdf(file_path: Path) -> list[LoadedDocument]:
    """逐頁載入 PDF，每頁產生一個 LoadedDocument。

    掃描式 PDF（內容是圖片而非文字圖層）會抽不到任何文字。
    本階段不支援 OCR，因此明確回報，而不是回傳空結果讓使用者困惑。

    Raises:
        DocumentLoadError: PDF 結構損毀而無法解析。
        EmptyDocumentError: 整份 PDF 抽不到任何文字（很可能是掃描式 PDF）。
    """
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(str(file_path))
        total_pages = len(reader.pages)
    except PdfReadError as exc:
        raise DocumentLoadError(
            f"PDF 檔案 '{file_path.name}' 無法解析，檔案可能已損毀"
        ) from exc
    except Exception as exc:  # pypdf 可能拋出多種底層例外
        raise DocumentLoadError(
            f"讀取 PDF '{file_path.name}' 時發生錯誤：{type(exc).__name__}"
        ) from exc

    documents: list[LoadedDocument] = []
    empty_pages: list[int] = []

    for index, page in enumerate(reader.pages):
        page_number = index + 1  # 頁碼從 1 開始，與 PDF 閱讀器顯示一致
        try:
            page_text = page.extract_text() or ""
        except Exception as exc:
            # 單頁失敗不應中斷整份文件，記錄後跳過
            logger.warning(
                "PDF 第 %d 頁文字抽取失敗（%s），已跳過該頁",
                page_number,
                type(exc).__name__,
            )
            empty_pages.append(page_number)
            continue

        if not page_text.strip():
            empty_pages.append(page_number)
            continue

        documents.append(
            LoadedDocument(
                text=page_text,
                source=file_path.name,
                file_name=file_path.name,
                file_type="pdf",
                page_number=page_number,
                paragraph_number=None,
                metadata={"total_pages": total_pages},
            )
        )

    if empty_pages:
        logger.info(
            "PDF '%s' 有 %d 頁無文字內容，已跳過（頁碼：%s）",
            file_path.name,
            len(empty_pages),
            empty_pages if len(empty_pages) <= 10 else f"{empty_pages[:10]}...",
        )

    if not documents:
        raise EmptyDocumentError(
            f"PDF '{file_path.name}' 共 {total_pages} 頁，但無法抽取任何文字。"
            "這可能是掃描式 PDF（內容為影像），本階段不支援 OCR。"
        )

    return documents


def load_docx(file_path: Path) -> list[LoadedDocument]:
    """載入 Word 文件，每個非空段落產生一個 LoadedDocument。

    paragraph_number 依「原始段落序號」計算（從 1 開始），空白段落雖然被跳過，
    但不會讓後續段落的編號前移 —— 這樣編號才能對應回使用者在 Word 中看到的位置。

    Raises:
        DocumentLoadError: DOCX 結構損毀而無法解析。
        EmptyDocumentError: 文件沒有任何非空段落。
    """
    import docx
    from docx.opc.exceptions import PackageNotFoundError

    try:
        document = docx.Document(str(file_path))
    except PackageNotFoundError as exc:
        raise DocumentLoadError(
            f"DOCX 檔案 '{file_path.name}' 無法開啟，檔案可能已損毀或不是有效的 Word 檔"
        ) from exc
    except Exception as exc:
        raise DocumentLoadError(
            f"讀取 DOCX '{file_path.name}' 時發生錯誤：{type(exc).__name__}"
        ) from exc

    documents: list[LoadedDocument] = []
    for index, paragraph in enumerate(document.paragraphs):
        paragraph_number = index + 1
        if not paragraph.text.strip():
            continue

        documents.append(
            LoadedDocument(
                text=paragraph.text,
                source=file_path.name,
                file_name=file_path.name,
                file_type="docx",
                page_number=None,
                paragraph_number=paragraph_number,
                metadata={},
            )
        )

    if not documents:
        raise EmptyDocumentError(f"DOCX '{file_path.name}' 沒有任何非空段落")

    return documents


def load_document(file_path: Path) -> list[LoadedDocument]:
    """統一入口：依副檔名分派到對應的載入器。

    Args:
        file_path: 檔案路徑（pathlib.Path）。

    Returns:
        文件單位清單。至少含一個元素，否則會拋出例外。

    Raises:
        DocumentLoadError: 檔案不存在、非一般檔案或讀取失敗。
        UnsupportedFileTypeError: 副檔名不支援。
        EmptyDocumentError: 檔案有效但抽不到文字。
    """
    _validate_file(file_path)

    suffix = file_path.suffix.lower()
    loaders = {
        ".txt": load_txt,
        ".pdf": load_pdf,
        ".docx": load_docx,
    }

    documents = loaders[suffix](file_path)
    logger.info(
        "已載入 '%s'（格式 %s），共 %d 個文件單位",
        file_path.name,
        suffix.lstrip("."),
        len(documents),
    )
    return documents
